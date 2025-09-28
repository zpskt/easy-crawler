#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
将爬取的文章内容导出为Word文档，保持图片与正文的相对位置
"""

import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from io import BytesIO
from urllib.parse import urlparse

def add_image_to_document(doc, image_info, max_width=6.0):
    """
    向Word文档添加图片
    
    Args:
        doc: Word文档对象
        image_info: 图片信息字典
        max_width: 图片最大宽度（英寸）
    """
    try:
        # 下载图片
        response = requests.get(image_info['url'], timeout=10)
        response.raise_for_status()
        
        # 将图片添加到文档
        image_stream = BytesIO(response.content)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加图片
        run = paragraph.add_run()
        picture = run.add_picture(image_stream, width=Inches(max_width))
        
        # 如果有alt文本，添加为图片说明
        if image_info.get('alt'):
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(f"图：{image_info['alt']}")
            caption_run.italic = True
            
    except Exception as e:
        # 如果图片下载失败，添加一个占位符段落
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"[图片不可用: {image_info['url']}]")
        run.italic = True

def export_article_to_word(article_data, output_path):
    """
    将单篇文章导出为Word文档
    
    Args:
        article_data: 文章数据字典
        output_path: 输出文件路径
    """
    # 创建Word文档
    doc = Document()
    
    # 添加标题
    title = doc.add_heading(article_data.get('title', '无标题'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加元数据信息
    metadata_para = doc.add_paragraph()
    metadata_run = metadata_para.add_run(
        f"来源: {article_data.get('url', '未知')}\n"
        f"发布时间: {article_data.get('publish_time', '未知')}\n"
        f"提取时间: {article_data.get('extraction_time', '未知')}"
    )
    metadata_run.font.size = doc.styles['Normal'].font.size
    metadata_run.font.italic = True
    
    # 处理正文内容
    content = article_data.get('content', '')
    images = {img.get('id', f"img_{i}"): img for i, img in enumerate(article_data.get('images', []))}
    
    # 按照内容中的占位符分割并插入内容和图片
    parts = content.split('[')
    doc.add_paragraph(parts[0])  # 添加第一部分文本
    
    for part in parts[1:]:
        if part.startswith('IMAGE_PLACEHOLDER_'):
            # 提取图片索引
            end_idx = part.find(']')
            if end_idx != -1:
                placeholder_content = part[:end_idx]
                remaining_content = part[end_idx+1:]
                
                # 提取图片ID
                try:
                    img_id = placeholder_content.split('_')[-1]
                    img_key = f"img_{img_id}"
                    if img_key in images:
                        add_image_to_document(doc, images[img_key])
                except:
                    pass
                
                # 添加剩余的文本内容
                if remaining_content.strip():
                    # 将文本按段落分割
                    paragraphs = remaining_content.strip().split('\n\n')
                    for paragraph_text in paragraphs:
                        if paragraph_text.strip():
                            doc.add_paragraph(paragraph_text.strip())
        else:
            # 这是普通的文本部分
            end_idx = part.find(']')
            if end_idx != -1:
                text_content = '[' + part
                if text_content.strip():
                    # 将文本按段落分割
                    paragraphs = text_content.strip().split('\n\n')
                    for paragraph_text in paragraphs:
                        if paragraph_text.strip():
                            doc.add_paragraph(paragraph_text.strip())
    
    # 保存文档
    doc.save(output_path)
    print(f"文档已保存至: {output_path}")

def export_articles_batch(json_file_path, output_dir):
    """
    批量导出文章为Word文档
    
    Args:
        json_file_path: 包含文章数据的JSON文件路径
        output_dir: 输出目录
    """
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)
    
    # 读取文章数据
    with open(json_file_path, 'r', encoding='utf-8') as f:
        articles = json.load(f)
    
    # 确保articles是列表格式
    if not isinstance(articles, list):
        articles = [articles]
    
    # 导出每篇文章
    for i, article in enumerate(articles):
        if 'error' in article:
            print(f"跳过文章 {i+1}: {article.get('error', '未知错误')}")
            continue
            
        # 生成文件名
        title = article.get('title', f'文章_{i+1}')
        # 清理文件名中的非法字符
        safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_title = safe_title.replace(' ', '_')
        output_filename = f"{safe_title[:50]}.docx"  # 限制文件名长度
        output_path = os.path.join(output_dir, output_filename)
        
        # 导出文章
        try:
            export_article_to_word(article, output_path)
        except Exception as e:
            print(f"导出文章失败 {title}: {e}")

def main():
    """
    主函数 - 示例用法
    """
    # 示例：导出单个JSON文件中的所有文章
    json_file_path = "/Users/zhangpeng/Desktop/zpskt/easy-crawler/outdir/extraction_result_1.json"
    output_dir = "../outdir/word_exports"
    
    if os.path.exists(json_file_path):
        export_articles_batch(json_file_path, output_dir)
    else:
        print(f"未找到文件: {json_file_path}")
        print("请先运行爬虫以生成数据文件")

if __name__ == "__main__":
    main()